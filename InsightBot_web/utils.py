import os
import time
import numpy as np

from pytube import Playlist
from youtube_transcript_api import YouTubeTranscriptApi

from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument

import pinecone
from pinecone import Pinecone
from groq import Groq

import ssl
import urllib.request
import os

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from dotenv import load_dotenv

# Global variables

# os.environ['PINECONE_API_KEY'] = os.getenv("PINECONE_API_KEY") # jatin
# os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY") # jatin
# os.environ.get('PINECONE_API_KEY')
# os.environ.get('GROQ_API_KEY')

def process_directory(directory_path):
    data = []
    
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            print(f"Processing file: {file_path}")
            
            # Skip empty files
            if os.path.getsize(file_path) == 0:
                print(f"Warning: The file {file_path} is empty and will be skipped.")
                continue
            
            try:
                # Process .txt files
                if file.endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        file_data = f.read()
                
                # Process .pdf files
                elif file.endswith(".pdf"):
                    reader = PdfReader(file_path)
                    file_data = ""
                    for page in reader.pages:
                        file_data += page.extract_text()
                
                # Process .docx files
                elif file.endswith(".docx"):
                    document = DocxDocument(file_path)
                    file_data = "\n".join([para.text for para in document.paragraphs])
                
                else:
                    print(f"Unsupported file format: {file}")
                    continue
                
                # Ensure valid content was extracted
                if not file_data.strip():
                    print(f"Warning: No content extracted from {file_path}. Skipping.")
                    continue
                
                # Append the processed data
                data.append({"File": file_path, "Data": file_data})
            
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")
    return data

def process_txt_file(file_path):
    print(f"Processing file: {file_path}")
    data = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            file_data = f.read()
            data.append({"File": file_path, "Data": file_data})
    except Exception as e:
                print(f"Error processing file {file_path}: {e}")
    return data

def delete_files_in_directory(directory_path):
    files = os.listdir(directory_path)
    for file in files:
        file_path = os.path.join(directory_path, file)
        if os.path.isfile(file_path):
            os.remove(file_path)
    print(f'All files from {directory_path} deleted successfully')

def prepare_data(documents):
    # Prepare the text for embedding
    document_data = []
    for document in documents:
        # Ensure 'Data' exists and is non-empty
        if 'Data' not in document or not document['Data']:
            print(f"Skipping document due to missing or empty 'Data': {document}")
            continue
        
        # Ensure the first element in 'Data' is a Document object
        if not isinstance(document['Data'][0], str):
            print(f"Skipping document due to invalid data type in 'Data': {document}")
            continue

        # Extract metadata and content
        document_source = document['File']
        document_content = document['Data']

        file_name = document_source.split("/")[-1]
        folder_names = document_source.split("/")[2:-1] if "/" in document_source else []

        doc = LangchainDocument(
            page_content=f"<Source>\n{document_source}\n</Source>\n\n<Content>\n{document_content}\n</Content>",
            metadata={
                "file_name": file_name,
                "parent_folder": folder_names[-1] if folder_names else "",
                "folder_names": folder_names
            }
        )
        document_data.append(doc)

    return document_data

def chunk_data(docs, chunk_size=1000,chunk_overlap=50):
    textsplitter=RecursiveCharacterTextSplitter(chunk_size=chunk_size,chunk_overlap=chunk_overlap)
    docs=textsplitter.split_documents(docs)
    return docs

class RAG:
    def __init__(self, model_name = "sentence-transformers/all-MiniLM-L6-v2"):
        load_dotenv()
        self.model_name = model_name
        self.embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        self.conversation_history=[]
        # Create an instance of the Pinecone
        self.pinecone = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
        # Create an instance of the Groq
        self.groq_client = Groq(api_key=os.environ.get('GROQ_API_KEY'))

    def get_huggingface_embeddings(self, text, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        model = SentenceTransformer(model_name)
        return model.encode(text)
    
    def initialize_pinecone(self, index_name: str):

        #Check if the index already exists; if not, raise an error or handle accordingly
        if index_name not in [index.name for index in self.pinecone.list_indexes().indexes]:
            raise ValueError(f"Index raise '{index_name}' does not exist. Please create it first.")

        # Connect to the specified index
        pinecone_index = self.pinecone.Index(index_name)

        return pinecone_index
    
    def upsert_vectorstore_to_pinecone(self, document_data, index_name, namespace):

        # Check if the namespace exists in the index
        pinecone_index = self.initialize_pinecone(index_name)

        # Check if the namespace exists by listing the namespaces (or by trying to query)
        namespaces = pinecone_index.describe_index_stats().get('namespaces', [])
        max_retries = 5
        wait_time = 2000
        if namespace in namespaces:
            print(f"Namespace '{namespace}' found. Deleting vector data...")
            pinecone_index.delete(namespace=namespace, delete_all=True)  # Initiates deletion

            # Polling to ensure deletion completes
            for attempt in range(max_retries):
                namespaces = pinecone_index.describe_index_stats().get('namespaces', [])
                if namespace not in namespaces:
                    print(f"Namespace '{namespace}' deletion confirmed.")
                    break
                time.sleep(wait_time)  # Wait before re-checking
            else:
                raise RuntimeError(f"Failed to delete namespace '{namespace}' after {max_retries} retries.")

        else:
            print(f"Namespace '{namespace}' does not exist. Proceeding with upsert.")

        # Create or replace the vector store
        vectorstore_from_documents = PineconeVectorStore.from_documents(
            document_data,
            self.embeddings,
            index_name=index_name,
            namespace=namespace
        )
        print(f"Vector store type: {type(vectorstore_from_documents)}")

        # Optionally, return the vector store if needed
        return vectorstore_from_documents

    def perform_rag(self, index_name, namespace, query):
        
        pinecone_index = self.initialize_pinecone(index_name)
        raw_query_embedding = self.get_huggingface_embeddings(query)

        query_embedding = np.array(raw_query_embedding)

        top_matches = pinecone_index.query(vector=query_embedding.tolist(), top_k=10, include_metadata=True, namespace=namespace)

        # Get the list of retrieved texts
        contexts = [item['metadata']['text'] for item in top_matches['matches']]

        augmented_query = "<CONTEXT>\n" + "\n\n-------\n\n".join(contexts[ : 10]) + "\n-------\n</CONTEXT>\n\n\n\nMY QUESTION:\n" + query
        
        conversations="\n".join([f"{msg['role'].upper()}:{msg['content']}" for msg in self.conversation_history])

        augmented_query += f"Conversation History:\n{conversations}\n\nMy Question:\n{query}"

        # Modify the prompt below as need to improve the response quality
        system_prompt = f'''
        You are a skilled expert in analyzing and understanding textual content from various sources, including YouTube video transcripts and document files.
        Your task is to answer any questions I have based on the provided text.
        If timestamps are present in seconds, convert them into a minutes:seconds format (e.g., 90 seconds becomes 1:30).
        Respond clearly and concisely with complete accuracy.
        '''

        res = self.groq_client.chat.completions.create(
            model="llama-3.1-70b-versatile", # llama-3.1-70b-versatile
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": augmented_query}
            ]
        )

        #Update conversations
        self.conversation_history.append({'role':'user','content':query})
        self.conversation_history.append({'role':'assistant','content':res.choices[0].message.content})

        return res.choices[0].message.content
    
    def reset_memory(self):
        """
        Clear the conversation history.
        """
        self.conversation_history = []

class YouTubeTranscriber:
    def __init__(self, url, transcript_language='en', output_dir='resources/transcripts'):
        """
        Initializes the YouTubeTranscriber with a URL (playlist or single video), transcript language, and output directory.

        :param url: URL of the YouTube playlist or single video
        :param transcript_language: Preferred language for transcripts (default is 'en' for English)
        :param output_dir: Directory where transcripts will be saved
        """
        self.url = url
        self.transcript_language = transcript_language
        self.output_dir = output_dir

        # Determine if the URL is a playlist or single video
        if 'playlist?list=' in url:
            self.is_playlist = True
            playlist_urls = self._get_playlist_urls(url)
            self.playlist_ids = []
            for pu in playlist_urls:
                self.playlist_ids.append(self._get_video_id(pu))
        elif 'watch?v=' in url or 'youtu.be/' in url:
            self.is_playlist = False
            self.video_id = self._get_video_id(url)
        else:
            raise ValueError("Invalid URL. Provide a valid YouTube playlist or video URL.")

        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        if self.is_playlist:
            print(f"Found {len(self.playlist)} videos in the playlist.")
    
    def _get_playlist_urls(self, playlist_url):
        """
        Extracts video URLs from a YouTube playlist.

        :param playlist_url: URL of the YouTube playlist
        :return: List of video URLs in the playlist
        """
        # Disable SSL verification
        ssl._create_default_https_context = ssl._create_unverified_context
        playlist_urls = Playlist(playlist_url)
        return list(playlist_urls)

    def _get_video_id(self, video_url):
        """
        Extracts the video ID from a YouTube video URL.

        :param video_url: URL of the YouTube video
        :return: Video ID as a string
        """
        if 'watch?v=' in video_url:
            return video_url.split('watch?v=')[1].split('&')[0]
        elif 'youtu.be/' in video_url:
            return video_url.split('youtu.be/')[1].split('?')[0]
        else:
            raise ValueError("Invalid video URL format.")

    def fetch_transcript(self, video_id):
        """
        Fetches the transcript for a single video in the specified language.

        :param video_id: YouTube video ID
        :return: Transcript as a list of dictionaries with 'start' and 'text' keys
        """
        try:
            # Fetch transcript with the specified language
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[self.transcript_language])
            return transcript
        except Exception as e:
            print(f"Could not retrieve transcript for video ID {video_id}: {e}")
            return None

    def save_transcript_to_file(self, video_id, transcript):
        """
        Saves the transcript of a video to a text file.

        :param video_id: YouTube video ID
        :param transcript: Transcript data to save
        """
        # Ensure the output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        file_path = os.path.join(self.output_dir, f"{video_id}_transcript.txt")
        with open(file_path, "w", encoding="utf-8") as file:
            for line in transcript:
                file.write(f"{line['start']}: {line['text']}\n")
        print(f"Transcript saved for video ID: {video_id}")

    def transcribe_playlist(self):
        """
        Processes each video in the playlist to fetch and save transcripts.
        """
        for video_id in self.playlist_ids:
            transcript = self.fetch_transcript(video_id)
            if transcript:
                self.save_transcript_to_file(video_id, transcript)

    def transcribe_single_video(self):
        """
        Fetches and saves the transcript for a single YouTube video.
        """
        # Fetch and save the transcript
        transcript = self.fetch_transcript(self.video_id)
        if transcript:
            self.save_transcript_to_file(self.video_id, transcript)

    def transcribe(self):
        """
        Determines if the URL is a playlist or single video and processes accordingly.
        """
        if self.is_playlist:
            self.transcribe_playlist()
        else:
            self.transcribe_single_video()
    
    def file_exists(self, video_id):
        """
        Checks if the transcript file for a given video ID already exists.

        :param video_id: YouTube video ID
        :return: True if file exists, False otherwise
        """
        file_path = os.path.join(self.output_dir, f"{video_id}_transcript.txt")
        return [os.path.isfile(file_path), file_path] # [True/False, path]
